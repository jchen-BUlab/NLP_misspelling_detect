#****************************************
# MIT License
# Copyright (c) 2026 Jinying Chen
#  
# author(s): Jinying Chen, Boston University Chobanian & Avedisian School of Medicine
# date: 2026-8-15
# ver: 1.0
# 
# This code was written to support model evaluation for the 2026 paper published 
# in JMIR Medical Informatics. 
# The code is for research use only, and is provided as it is.
# 
# See LICENSE in the project root for license terms.

import argparse
import json
import os
from pathlib import Path

import pandas as pd

import evaluate_gpt_api as base_eval
from create_dataset_for_gpt_api_expt import normalize_text


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument("--config", required=True, help="Path to a JSON config file.")
    return parser.parse_args()


def load_config(config_path):
    with open(config_path, "r", encoding="utf-8") as handle:
        config = json.load(handle)

    if "model" not in config or not str(config["model"]).strip():
        raise ValueError("Config file must contain a non-empty 'model' field")
    if "trial_folders" not in config or not config["trial_folders"]:
        raise ValueError("Config file must contain a non-empty 'trial_folders' list")

    config["model"] = str(config["model"]).strip()
    config["trial_folders"] = [str(folder).strip() for folder in config["trial_folders"] if str(folder).strip()]
    if not config["trial_folders"]:
        raise ValueError("Config file must contain at least one valid trial folder")

    output_folder = config.get("output_folder")
    if output_folder:
        config["output_folder"] = str(output_folder).strip()
    else:
        config["output_folder"] = config["trial_folders"][0]

    source_filter = config.get("source_filter", "RxNorm")
    config["source_filter"] = str(source_filter).strip() if str(source_filter).strip() else "RxNorm"

    return config


def collect_trial_files(trial_folders):
    files_by_name = {}
    for folder in trial_folders:
        folder_path = Path(folder)
        for file_path in base_eval.list_scannable_output_files(str(folder_path)):
            filename = os.path.basename(file_path)
            files_by_name.setdefault(filename, []).append(file_path)
    return files_by_name


def infer_prediction_column(df, file_path):
    prediction_column, _ = base_eval.infer_prediction_columns(df, file_path)
    return prediction_column


def infer_prediction_and_score_columns(df, file_path):
    prediction_column, probability_column = base_eval.infer_prediction_columns(df, file_path)
    score_column = probability_column if probability_column else prediction_column
    return prediction_column, score_column


def build_trial_predictions_df(file_paths, source_filter):
    trial_frames = []
    trial_labels = []
    target_seen = False

    for trial_index, file_path in enumerate(sorted(file_paths), start=1):
        df = pd.read_csv(file_path)
        df = base_eval.filter_by_source(df, source_filter)
        if df.empty:
            print(f"Warning: no rows matched source_filter={source_filter} in {file_path}; skipping this trial")
            continue

        prediction_column = infer_prediction_column(df, file_path)
        trial_label = f"trial_{trial_index}"
        trial_labels.append(trial_label)

        trial_df = df.copy()
        trial_df["input_norm"] = trial_df["input"].apply(normalize_text)
        trial_df = trial_df.drop_duplicates(subset=["input_norm"])

        selected_columns = ["input_norm", "input", prediction_column]
        rename_map = {prediction_column: trial_label}
        if "target" in trial_df.columns and not target_seen:
            selected_columns.append("target")
            target_seen = True

        trial_df = trial_df[selected_columns].rename(columns=rename_map)
        trial_frames.append(trial_df)

    if not trial_frames:
        return pd.DataFrame(), []

    merged_df = trial_frames[0]
    for trial_df in trial_frames[1:]:
        if "target" in trial_df.columns:
            trial_df = trial_df.drop(columns=["target"])
        if "input" in trial_df.columns:
            trial_df = trial_df.drop(columns=["input"])
        merged_df = merged_df.merge(trial_df, on="input_norm", how="outer")

    if "input" not in merged_df.columns:
        merged_df["input"] = merged_df["input_norm"]
    else:
        merged_df["input"] = merged_df["input"].where(
            merged_df["input"].notna(),
            merged_df["input_norm"],
        )

    return merged_df, trial_labels


def calculate_trial_metrics(file_paths, source_filter):
    trial_metrics = []

    for trial_index, file_path in enumerate(sorted(file_paths), start=1):
        df = pd.read_csv(file_path)
        df = base_eval.filter_by_source(df, source_filter)
        if df.empty:
            print(f"Warning: no rows matched source_filter={source_filter} in {file_path}; skipping trial metrics")
            continue

        prediction_column, score_column = infer_prediction_and_score_columns(df, file_path)
        metrics = base_eval.summarize_metrics(df, prediction_column, score_column)
        metrics["trial_name"] = f"trial_{trial_index}"
        metrics["file_path"] = file_path
        trial_metrics.append(metrics)

    return trial_metrics


def calculate_consistency(trial_predictions_df, trial_labels):
    if trial_predictions_df.empty or not trial_labels:
        return pd.DataFrame()

    detail_df = trial_predictions_df.copy()
    detail_df["num_trials_used"] = detail_df[trial_labels].notna().sum(axis=1).astype(int)
    detail_df["predict_1_count"] = detail_df[trial_labels].eq(1).sum(axis=1).astype(int)
    detail_df["predict_0_count"] = detail_df[trial_labels].eq(0).sum(axis=1).astype(int)

    denominator = detail_df["num_trials_used"].replace(0, pd.NA)
    detail_df["predict_1_percent"] = detail_df["predict_1_count"] / denominator
    detail_df["predict_0_percent"] = detail_df["predict_0_count"] / denominator
    detail_df["consistency_score"] = detail_df[["predict_1_percent", "predict_0_percent"]].max(axis=1)
    return detail_df


def build_summary_tsv_filename(model, source_filter):
    suffix = base_eval.sanitize_config_name(source_filter)
    return f"{base_eval.sanitize_config_name(model)}_stability_summary_{suffix}.tsv"


def build_summary_docx_filename(model, source_filter):
    suffix = base_eval.sanitize_config_name(source_filter)
    return f"{base_eval.sanitize_config_name(model)}_stability_summary_{suffix}.docx"


def derive_prompt_type(setting_file):
    metadata = base_eval.parse_filename_metadata(setting_file)
    prompt_type = str(metadata.get("prompt_style", "")).strip()
    return prompt_type if prompt_type else "unknown"


def build_summary_report_df(summary_df):
    report_columns = [
        "model",
        "prompt_type",
        "overall_consistency_score",
        "mean_precision",
        "sd_precision",
        "mean_recall",
        "sd_recall",
        "mean_f1",
        "sd_f1",
        "mean_accuracy",
        "sd_accuracy",
        "mean_specificity",
        "sd_specificity",
        "mean_Brier score",
        "sd_Brier score",
        "mean_ROC AUC",
        "sd_ROC AUC",
        "mean_PR AUC",
        "sd_PR AUC",
        "num_trials",
        "num_terms",
        "setting_file",
    ]

    if summary_df.empty:
        return pd.DataFrame(columns=report_columns)

    report_df = summary_df.copy()
    report_df["prompt_type"] = report_df["setting_file"].apply(derive_prompt_type)
    report_df["model"] = report_df["setting_file"].apply(lambda value: Path(value).stem)
    report_df = report_df[report_columns].copy()
    return report_df.sort_values(["prompt_type", "model"], ascending=[True, True]).reset_index(drop=True)


def format_report_value(value):
    if pd.isna(value):
        return "NA"
    if isinstance(value, (int, float)):
        return f"{float(value):.3f}"
    return str(value)


def format_mean_sd(mean_value, sd_value):
    if pd.isna(mean_value) and pd.isna(sd_value):
        return "NA"
    if pd.isna(mean_value):
        return f"NA +- {format_report_value(sd_value)}"
    if pd.isna(sd_value):
        return f"{format_report_value(mean_value)} +- NA"
    return f"{format_report_value(mean_value)} +- {format_report_value(sd_value)}"


def write_summary_tsv(summary_df, output_path):
    build_summary_report_df(summary_df).to_csv(output_path, sep="\t", index=False)


def write_summary_docx(summary_df, output_path):
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to write the stability Word report. "
            "Please install it in this environment."
        ) from exc

    report_df = build_summary_report_df(summary_df)
    report_columns = [
        "model",
        "prompt_type",
        "overall_consistency_score",
        "precision",
        "recall",
        "f1",
        "accuracy",
        "specificity",
        "Brier score",
        "ROC AUC",
        "PR AUC",
        "num_trials",
        "num_terms",
    ]

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    document.add_heading("Model Stability Summary", level=1)

    if report_df.empty:
        document.add_paragraph("No stability summary rows were generated.")
        document.save(output_path)
        return

    table = document.add_table(rows=1, cols=len(report_columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, column_name in enumerate(report_columns):
        header_cells[index].text = column_name

    for _, row in report_df.iterrows():
        cells = table.add_row().cells
        row_values = {
            "model": row["model"],
            "prompt_type": row["prompt_type"],
            "overall_consistency_score": format_report_value(row["overall_consistency_score"]),
            "precision": format_mean_sd(row["mean_precision"], row["sd_precision"]),
            "recall": format_mean_sd(row["mean_recall"], row["sd_recall"]),
            "f1": format_mean_sd(row["mean_f1"], row["sd_f1"]),
            "accuracy": format_mean_sd(row["mean_accuracy"], row["sd_accuracy"]),
            "specificity": format_mean_sd(row["mean_specificity"], row["sd_specificity"]),
            "Brier score": format_mean_sd(row["mean_Brier score"], row["sd_Brier score"]),
            "ROC AUC": format_mean_sd(row["mean_ROC AUC"], row["sd_ROC AUC"]),
            "PR AUC": format_mean_sd(row["mean_PR AUC"], row["sd_PR AUC"]),
            "num_trials": format_report_value(row["num_trials"]),
            "num_terms": format_report_value(row["num_terms"]),
        }
        for index, column_name in enumerate(report_columns):
            cells[index].text = str(row_values[column_name])

    document.save(output_path)


def summarize_setting(model, filename, detail_df, trial_count):
    summary_row = {
        "model": model,
        "setting_file": filename,
        "num_trials": trial_count,
        "num_terms": 0 if detail_df.empty else int(len(detail_df)),
        "overall_consistency_score": pd.NA if detail_df.empty else float(detail_df["consistency_score"].mean()),
        "mean_predict_1_percent": pd.NA if detail_df.empty else float(detail_df["predict_1_percent"].mean()),
        "mean_predict_0_percent": pd.NA if detail_df.empty else float(detail_df["predict_0_percent"].mean()),
    }
    return summary_row


def add_metric_summary(summary_row, trial_metrics):
    metric_names = ["precision", "recall", "f1", "accuracy", "specificity", "Brier score", "ROC AUC", "PR AUC"]
    if not trial_metrics:
        for metric_name in metric_names:
            summary_row[f"mean_{metric_name}"] = pd.NA
            summary_row[f"sd_{metric_name}"] = pd.NA
        return summary_row

    trial_metrics_df = pd.DataFrame(trial_metrics)
    for metric_name in metric_names:
        metric_series = pd.to_numeric(trial_metrics_df[metric_name], errors="coerce").dropna()
        if metric_series.empty:
            summary_row[f"mean_{metric_name}"] = pd.NA
            summary_row[f"sd_{metric_name}"] = pd.NA
            continue
        summary_row[f"mean_{metric_name}"] = float(metric_series.mean())
        summary_row[f"sd_{metric_name}"] = float(metric_series.std(ddof=0))
    return summary_row


def main():
    args = parse_args()
    config = load_config(args.config)

    model = config["model"]
    trial_folders = config["trial_folders"]
    output_folder = Path(config["output_folder"])
    source_filter = config["source_filter"]
    output_folder.mkdir(parents=True, exist_ok=True)

    files_by_name = collect_trial_files(trial_folders)
    repeated_files = {
        filename: paths
        for filename, paths in files_by_name.items()
        if len(paths) >= 2
    }

    summary_rows = []
    detail_rows = []

    for filename, file_paths in sorted(repeated_files.items()):
        trial_predictions_df, trial_labels = build_trial_predictions_df(file_paths, source_filter)
        detail_df = calculate_consistency(trial_predictions_df, trial_labels)
        trial_metrics = calculate_trial_metrics(file_paths, source_filter)
        summary_row = summarize_setting(model, filename, detail_df, len(trial_labels))
        summary_row = add_metric_summary(summary_row, trial_metrics)
        summary_row["source_filter"] = source_filter
        summary_rows.append(summary_row)

        if not detail_df.empty:
            detail_df = detail_df.copy()
            detail_df.insert(0, "model", model)
            detail_df.insert(1, "setting_file", filename)
            detail_df.insert(2, "num_trials_total", len(trial_labels))
            detail_df.insert(3, "source_filter", source_filter)
            detail_rows.append(detail_df)

    summary_df = pd.DataFrame(summary_rows).sort_values(
        ["overall_consistency_score", "setting_file"],
        ascending=[False, True],
        na_position="last",
    )
    detail_output_df = pd.concat(detail_rows, ignore_index=True) if detail_rows else pd.DataFrame()

    suffix = base_eval.sanitize_config_name(source_filter)
    summary_output_path = output_folder / f"{base_eval.sanitize_config_name(model)}_stability_summary_{suffix}.csv"
    detail_output_path = output_folder / f"{base_eval.sanitize_config_name(model)}_stability_detail_{suffix}.csv"
    summary_tsv_output_path = output_folder / build_summary_tsv_filename(model, source_filter)
    summary_docx_output_path = output_folder / build_summary_docx_filename(model, source_filter)

    summary_df.to_csv(summary_output_path, index=False)
    detail_output_df.to_csv(detail_output_path, index=False)
    write_summary_tsv(summary_df, summary_tsv_output_path)
    write_summary_docx(summary_df, summary_docx_output_path)

    print(f"Wrote {summary_output_path}")
    print(f"Wrote {detail_output_path}")
    print(f"Wrote {summary_tsv_output_path}")
    print(f"Wrote {summary_docx_output_path}")


if __name__ == "__main__":
    main()
